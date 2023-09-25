import random #IMPORT DE BIBLIOTECA DE NUMEROS ALEATORIOS
import  firebirdsql #IMPORT DE BIBLIOTECA PARA CONEXÃO
import docx #IMPORT DE BIBLIOTECA PARA MANIPULAÇÃO DE WORD
import os
import smtplib
from docx import Document
from datetime import datetime, timedelta
from email.message import EmailMessage


email = "testeprojetotg@gmail.com"
senha = "vneogzuuyfdsdtbl"
msg =  EmailMessage()
msg['Subject'] = 'Enviando Email com Python'
msg['From'] =  "testeprojetotg@gmail.com"
emailEnvio = input('Digite o email de envio: ')
msg['To'] = emailEnvio
msg.set_content("Segue a Escala Automatizada")


#========CONEXÃO======COM======BANCO=================
con = firebirdsql.Connection(user="SYSDBA",password="masterkey",database="A:/Biblioteca/Documents/Embarcadero/Studio/Projects/ProjetoExp/DB/dbteste.fdb", host="localhost",)
#CONECTANDO COM O BANCO
cur = con.cursor()
#EXECUTAR COMANDOS NO BANCO


#CRIAÇAÕ DE DOCUMENTO WORD
doc = docx.Document()
doc.add_heading('Modelo Escala', 0)

def sortear_numeros(qtd_numeros, numeros_disponiveis):
    if len(numeros_disponiveis) < qtd_numeros:
        numeros_disponiveis.extend(numeros_gerais)
        random.shuffle(numeros_disponiveis)

    numeros_sorteados = []
    for _ in range(qtd_numeros):
        numero = numeros_disponiveis.pop()
        while numero in monitores_gerais:
            numero = numeros_disponiveis.pop()
        numeros_sorteados.append(numero)

    return numeros_sorteados

def sortear_monitor(qtd_numeros, monitores_disponiveis):
    monitores_sorteados = []
    while len(monitores_sorteados) < qtd_numeros:
        if len(monitores_disponiveis) < qtd_numeros:
            monitores_disponiveis.extend(monitores_gerais)
            random.shuffle(monitores_disponiveis)

        monitor_sorteado = monitores_disponiveis.pop()
        while monitor_sorteado in monitores_sorteados:
            monitor_sorteado = monitores_disponiveis.pop()

        monitores_sorteados.append(monitor_sorteado)

    return monitores_sorteados

data = datetime.now()
ano = int(input('Digite o ano desejado: '))
mes = int(input('Digite o mês desejado: '))
dia = int(input('Digite o dia desejado: '))
data_limite = datetime(ano, mes, dia)
numeros_gerais = list(range(1, 101))
monitores_gerais = [1, 6, 18, 24, 36, 37, 45, 48, 49, 53, 66, 74, 76, 77, 78, 87, 89, 100]
numeros_sorteados_geral = list(numeros_gerais)  # a lista geral possui todos os números no início
numeros_monitores_geral = list(monitores_gerais)
random.shuffle(numeros_sorteados_geral)
random.shuffle(numeros_monitores_geral)

#armazenar os monitores sorteados e disponíveis
monitores_sorteados = []
monitores_disponiveis = list(numeros_monitores_geral)
ciclos = 0


while data < data_limite:
    data = data + timedelta(days=1)
    quantidade = 3
    quantidade_monitores = 1
    numeros_sorteados = sortear_numeros(quantidade, numeros_sorteados_geral)
    monitores_sorteio_atual = sortear_monitor(quantidade_monitores, monitores_disponiveis)

        #atualização a lista de monitores sorteados
    monitores_sorteados.extend(monitores_sorteio_atual)

        #print("Monitor:", monitores_sorteio_atual, "Atiradores:", numeros_sorteados, "Data:", data.date(), "\n")

    InserirEscala = [
        (monitores_sorteio_atual[0], numeros_sorteados[0],numeros_sorteados[1],numeros_sorteados[2], data.date())
        ]
    cur.executemany("insert into escala (id_monitor, nmr_atirador1, nmr_atirador2, nmr_atirador3, data) values (?,?,?,?,?)", InserirEscala)
    con.commit()

    #
        
    listaData =[(data.date())]
    #QUERY 1 - MONITOR
    cur.execute("select  monitores.nome as monitor, monitores.numero as nmr, escala.data as data from monitores, escala where monitores.numero = escala.id_monitor and escala.data  = ? " , listaData)
    r= cur.fetchone()
    NomeMonitor, nmrMonitor, dataEscala = r
    print(NomeMonitor, nmrMonitor, dataEscala)
    
    # QUERY2- ATIRADOR1
    cur.execute("select atiradores.nome as atirador1, atiradores.numero as numero,escala.data as data  from atiradores, escala where atiradores.numero = escala.nmr_atirador1 and data = ?", listaData)
    At1 = cur.fetchone()
    NomeAt1, nmrAt1, dataEscala = At1
    print(NomeAt1, nmrAt1, dataEscala)
    
    # QUERY3- ATIRADOR2
    cur.execute("select atiradores.nome as atirado2, atiradores.numero as numero,escala.data as data  from atiradores, escala where atiradores.numero = escala.nmr_atirador2 and data = ?", listaData)
    At2 = cur.fetchone()
    NomeAt2, nmrAt2, dataEscala = At2
    print(NomeAt2, nmrAt2, dataEscala)
        
        # QUERY4- ATIRADOR1
    cur.execute("select atiradores.nome as atirador3, atiradores.numero as numero,escala.data as data  from atiradores, escala where atiradores.numero = escala.nmr_atirador3 and data = ?", listaData)
    At3 = cur.fetchone()
    NomeAt3, nmrAt3, dataEscala = At3
    print(NomeAt3, nmrAt3, dataEscala)
        
        
        
        #ESCALA - CRIAR - WORD
        
        
    localCmt ='Cmt Gda / Permanencia'
    localSnt = 'P1'
    dados =(
    (localCmt, nmrMonitor,NomeMonitor),
    (localSnt, nmrAt1, NomeAt1),
    (localSnt, nmrAt2, NomeAt2),
    (localSnt, nmrAt3, NomeAt3)
    )
    doc.add_paragraph(f'Data: {data.date()}')
    table = doc.add_table(rows=1, cols=3)
    table.Style = 'Medium shadding 5 Accent 3'
    row = table.rows[0].cells
    row[0].text = 'Local'
    row[1].text = 'Numero'
    row[2].text = 'Atirador'

    for Local, Numero, Atirador in dados:
        row = table.add_row().cells
        row[0].text = Local
        row[1].text = str(Numero)
        row[2].text = Atirador
        doc.add_paragraph()

        #verificação se existem números repetidos
    numeros_unicos = set(numeros_sorteados_geral)
    if len(numeros_unicos) != len(numeros_sorteados_geral):
        print("Ha numeros repetidos na lista de numeros sorteados. \n")
    else:
        print("Nao ha numeros repetidos na lista de numeros sorteados.\n")

        #verificação se um ciclo completo foi concluído
    if len(monitores_disponiveis) == 0:
        monitores_disponiveis = list(numeros_monitores_geral)  # reiniciar lista de monitores disponíveis
        ciclos += 1

doc.save(f"EscalaAutomatizada1.docx")

print("Lista de atiradores disponiveis:", list(set(numeros_sorteados_geral) - set(monitores_gerais)))
print("Lista de monitores disponiveis:", monitores_disponiveis)


with open('EscalaAutomatizada1.docx','rb') as content_file:
    content = content_file.read()
    msg.add_attachment(content, maintype='application', subtype='docx', filename='EscalaAtualizada1.docx')

with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
    smtp.login(email, senha)
    smtp.send_message(msg)

